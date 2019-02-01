# coding:utf-8
""" word替换处理 """

from docx import Document
import os

class DocxHandle(object):

    def __init__(self, doc_file,  data,  rule, out=print,  err=print):
        """ DocxHandle.
            Args: 
                doc_file: template doc file.
                data: 数据
                rule: 规则
                out: 输出重定向
        """
        self.document = Document(doc_file)  # 打开文件docx文件
        self.rule = rule
        self.data = data
        self.log = out
        self.err = err
    
    def save(self, title):
        """ 保存文档 """
        # 新建目录
        img_dir = './docxdata/'
        if not os.path.exists(img_dir):
            os.mkdir(img_dir)
        file_name = f"{img_dir}{title}.docx"
        self.document.save(file_name) 
        self.log(f"存储文件：{img_dir}{title}.docx")
        return file_name
            
    def test_rule_template(self):
        for x,y in self.rule.items():
            value = self.data[y]
            self.log(f"{x},[{value}]")
            # 判断是否有未匹配的模板变量
            if not self.replace_text(x, value, self.document, debug=True):
                self.err(f"{x}模板未找到，请重试")
                return False
        return True

    def run(self):
        """ 运行，检测模板变量替换 """
        for x,y in self.rule.items():
            value = self.data[y]
            self.log(f"{x},[{value}]")
            self.replace_text(x, value, self.document)
        return self.document

    def replace_text(self, old_text, new_text, file, debug=False):
        """替换文本， 传入参数： 旧字符串, 新字符串, 文件对象，是否调试"""
        # 标记是否替换
        flag = False
        # 遍历文件对象
        for f in file.paragraphs:
            # 如果 旧字符串 在 某个段落 中
            if old_text in f.text:
                if debug:
                    self.log(f"替换前===>{f.text}")
                # 遍历 段落 生成 i
                for i in f.runs:
                    # 如果 旧字符串 在 i 中
                    if debug:
                        self.log(i.text)
                    if old_text in i.text:
                        # 替换 i.text 内文本资源
                        i.text = i.text.replace(old_text, new_text)
                        flag = True
                if debug:
                    self.log(f"替换后===>{f.text} {flag}")
        
        return flag
       

class MergeDocx(object):
    """ 合并docx """
    
    def __init__(self, log=print):
        self.document = Document()
        self.log=log
    
    def append(self,  new_doc_obj):
        """ 追加文档，new_doc_obj为Document对象 """
#        for paragraph in new_doc_obj.paragraphs:
#            text = paragraph.text
#            self.document.add_paragraph(text)
#            
        for element in new_doc_obj.element.body:
            self.document.element.body.append(element)
        else:
            #self.document.add_page_break()
            pass
        
    def save(self, title):
        """ 保存文档 """
        # 新建目录
        img_dir = './docxdata/'
        if not os.path.exists(img_dir):
            os.mkdir(img_dir)
        file_name = f"{img_dir}{title}.docx"
        self.document.save(file_name) 
        self.log(f"存储文件：{img_dir}{title}.docx")
        return file_name
    
    def __del__(self):
        self.save("result")
